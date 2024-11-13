import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os
import pyperclip  # Pour copier dans le presse-papiers

def create_creation_ticket_installation_tab(tab_creation_ticket_installation):
    # Créer un canvas pour le scroll
    canvas = tk.Canvas(tab_creation_ticket_installation)
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Ajouter un scrollbar vertical lié au canvas
    scrollbar = tk.Scrollbar(tab_creation_ticket_installation, orient=tk.VERTICAL, command=canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    # Créer un frame à l'intérieur du canvas pour y ajouter les widgets
    scrollable_frame = tk.Frame(canvas)

    # Configurer le canvas pour utiliser le scrollbar
    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(
            scrollregion=canvas.bbox("all")
        )
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

    # S'assurer que le canvas réagit à la molette de la souris
    scrollable_frame.bind_all("<MouseWheel>", lambda event: canvas.yview_scroll(int(-1*(event.delta/120)), "units"))

    # Largeur uniforme pour tous les champs
    field_width = 60  # Largeur des champs
    large_text_height = 6  # Hauteur des Text widgets pour les sections comme Description et Actions

    # Configuration des colonnes pour une meilleure gestion de l'espace
    scrollable_frame.columnconfigure(0, weight=1)
    scrollable_frame.columnconfigure(1, weight=3)

    # Widgets pour les inputs à gauche
    label_titre = tk.Label(scrollable_frame, text="Titre du ticket :")
    label_titre.grid(row=0, column=0, padx=10, pady=5, sticky="e")
    entry_titre = tk.Entry(scrollable_frame, width=field_width)
    entry_titre.grid(row=0, column=1, padx=10, pady=5, sticky="w")

    label_date = tk.Label(scrollable_frame, text="Date et heure :")
    label_date.grid(row=1, column=0, padx=10, pady=5, sticky="e")
    entry_date = tk.Entry(scrollable_frame, width=field_width)
    entry_date.insert(0, datetime.now().strftime("%Y-%m-%d %H:%M"))
    entry_date.grid(row=1, column=1, padx=10, pady=5, sticky="w")

    label_client = tk.Label(scrollable_frame, text="Client :")
    label_client.grid(row=2, column=0, padx=10, pady=5, sticky="e")
    entry_client = tk.Entry(scrollable_frame, width=field_width)
    entry_client.grid(row=2, column=1, padx=10, pady=5, sticky="w")

    # Champs supplémentaires
    label_version = tk.Label(scrollable_frame, text="Version installée :")
    label_version.grid(row=3, column=0, padx=10, pady=5, sticky="e")
    entry_version = tk.Entry(scrollable_frame, width=field_width)
    entry_version.grid(row=3, column=1, padx=10, pady=5, sticky="w")

    label_lecteurs = tk.Label(scrollable_frame, text="Type lecteurs installés :")
    label_lecteurs.grid(row=4, column=0, padx=10, pady=5, sticky="e")
    entry_lecteurs = tk.Entry(scrollable_frame, width=field_width)
    entry_lecteurs.grid(row=4, column=1, padx=10, pady=5, sticky="w")

    label_fournisseurLecteurs = tk.Label(scrollable_frame, text="Fournisseur lecteurs installés :")
    label_fournisseurLecteurs.grid(row=4, column=0, padx=10, pady=5, sticky="e")
    entry_fournisseurLecteurs = tk.Entry(scrollable_frame, width=field_width)
    entry_fournisseurLecteurs.grid(row=4, column=1, padx=10, pady=5, sticky="w")

    label_sauvegarde = tk.Label(scrollable_frame, text="Présence sauvegarde :")
    label_sauvegarde.grid(row=5, column=0, padx=10, pady=5, sticky="e")
    entry_sauvegarde = tk.Entry(scrollable_frame, width=field_width)
    entry_sauvegarde.grid(row=5, column=1, padx=10, pady=5, sticky="w")

    label_postes = tk.Label(scrollable_frame, text="Nom des postes :")
    label_postes.grid(row=6, column=0, padx=10, pady=5, sticky="ne")
    listbox_postes = tk.Listbox(scrollable_frame, height=4, width=field_width)
    listbox_postes.grid(row=6, column=1, padx=10, pady=5, sticky="w")

    btn_ajouter_poste = tk.Button(scrollable_frame, text="Ajouter Poste", command=lambda: ajouter_poste(listbox_postes))
    btn_ajouter_poste.grid(row=7, column=1, padx=10, pady=5, sticky="w")

    label_poste_serveur = tk.Label(scrollable_frame, text="Poste serveur + type poste :")
    label_poste_serveur.grid(row=8, column=0, padx=10, pady=5, sticky="e")
    entry_poste_serveur = tk.Entry(scrollable_frame, width=field_width)
    entry_poste_serveur.grid(row=8, column=1, padx=10, pady=5, sticky="w")

    # Description et autres sections avec hauteur ajustée
    label_description = tk.Label(scrollable_frame, text="Description de l'installation :")
    label_description.grid(row=9, column=0, padx=10, pady=5, sticky="ne")
    text_description = tk.Text(scrollable_frame, wrap='word', height=large_text_height, width=field_width)
    text_description.grid(row=9, column=1, padx=10, pady=5, sticky="w")

    label_actions = tk.Label(scrollable_frame, text="Actions à suivre :")
    label_actions.grid(row=10, column=0, padx=10, pady=5, sticky="ne")
    text_actions = tk.Text(scrollable_frame, wrap='word', height=large_text_height, width=field_width)
    text_actions.grid(row=10, column=1, padx=10, pady=5, sticky="w")

    label_images = tk.Label(scrollable_frame, text="Images (Screenshots) :")
    label_images.grid(row=11, column=0, padx=10, pady=5, sticky="ne")
    listbox_images = tk.Listbox(scrollable_frame, height=4, width=field_width)
    listbox_images.grid(row=11, column=1, padx=10, pady=5, sticky="w")

    btn_ajouter_image = tk.Button(scrollable_frame, text="Ajouter Image", command=lambda: ajouter_image(listbox_images))
    btn_ajouter_image.grid(row=12, column=1, padx=10, pady=5, sticky="w")

    # Visionneuse à droite
    label_visionneuse = tk.Label(scrollable_frame, text="Aperçu du Ticket d'Installation :")
    label_visionneuse.grid(row=0, column=2, padx=10, pady=10, sticky="w")
    text_visionneuse = tk.Text(scrollable_frame, wrap='word', height=40, width=50)
    text_visionneuse.grid(row=1, column=2, rowspan=12, padx=10, pady=5, sticky="nsew")

    # Définir le tag pour le texte en gras dans la visionneuse
    text_visionneuse.tag_configure("bold", font=("Helvetica", 10, "bold"))

    # Bouton pour copier le contenu de la visionneuse
    btn_copier = tk.Button(scrollable_frame, text="Copier", command=lambda: copier_texte_visionneuse(text_visionneuse))
    btn_copier.grid(row=13, column=2, padx=10, pady=5, sticky="e")

    def mettre_a_jour_visionneuse(event=None):
        text_visionneuse.config(state=tk.NORMAL)
        text_visionneuse.delete(1.0, tk.END)
        titre = entry_titre.get().strip()
        date_heure = entry_date.get().strip()
        client = entry_client.get().strip()
        version = entry_version.get().strip()
        lecteurs = entry_lecteurs.get().strip()
        fournisseurLecteurs = entry_fournisseurLecteurs.get().strip()
        sauvegarde = entry_sauvegarde.get().strip()
        postes = listbox_postes.get(0, tk.END)
        poste_serveur = entry_poste_serveur.get().strip()
        description = text_description.get("1.0", tk.END).strip()
        actions = text_actions.get("1.0", tk.END).strip()

        # Mise en forme des titres en gras à l'aide de tags
        text_visionneuse.insert(tk.END, "Titre : ", "bold")
        text_visionneuse.insert(tk.END, f"{titre}\n")
        text_visionneuse.insert(tk.END, "Date et Heure : ", "bold")
        text_visionneuse.insert(tk.END, f"{date_heure}\n")
        text_visionneuse.insert(tk.END, "Client : ", "bold")
        text_visionneuse.insert(tk.END, f"{client}\n")
        text_visionneuse.insert(tk.END, "Version installée : ", "bold")
        text_visionneuse.insert(tk.END, f"{version}\n")
        text_visionneuse.insert(tk.END, "Type lecteurs installés : ", "bold")
        text_visionneuse.insert(tk.END, f"{lecteurs}\n")
        text_visionneuse.insert(tk.END, "Fournisseur des lecteurs installés : ", "bold")
        text_visionneuse.insert(tk.END, f"{fournisseurLecteurs}\n")
        text_visionneuse.insert(tk.END, "Présence sauvegarde : ", "bold")
        text_visionneuse.insert(tk.END, f"{sauvegarde}\n")
        text_visionneuse.insert(tk.END, "Nom des postes : ", "bold")
        text_visionneuse.insert(tk.END, f"{', '.join(postes)}\n")
        text_visionneuse.insert(tk.END, "Poste serveur + type poste : ", "bold")
        text_visionneuse.insert(tk.END, f"{poste_serveur}\n\n")
        text_visionneuse.insert(tk.END, "Description :\n", "bold")
        text_visionneuse.insert(tk.END, f"{description}\n\n")
        text_visionneuse.insert(tk.END, "Actions à suivre :\n", "bold")
        text_visionneuse.insert(tk.END, f"{actions}")

        text_visionneuse.config(state=tk.DISABLED)

    # Lier les champs d'entrée à la fonction de mise à jour de la visionneuse
    entry_titre.bind("<KeyRelease>", mettre_a_jour_visionneuse)
    entry_date.bind("<KeyRelease>", mettre_a_jour_visionneuse)
    entry_client.bind("<KeyRelease>", mettre_a_jour_visionneuse)
    entry_version.bind("<KeyRelease>", mettre_a_jour_visionneuse)
    entry_lecteurs.bind("<KeyRelease>", mettre_a_jour_visionneuse)
    entry_sauvegarde.bind("<KeyRelease>", mettre_a_jour_visionneuse)
    listbox_postes.bind("<KeyRelease>", mettre_a_jour_visionneuse)
    entry_poste_serveur.bind("<KeyRelease>", mettre_a_jour_visionneuse)
    text_description.bind("<KeyRelease>", mettre_a_jour_visionneuse)
    text_actions.bind("<KeyRelease>", mettre_a_jour_visionneuse)

    def ajouter_poste(listbox_postes):
        poste = simpledialog.askstring("Nom du Poste", "Entrez le nom du poste :")
        if poste:
            listbox_postes.insert(tk.END, poste)
        mettre_a_jour_visionneuse()

    def ajouter_image(listbox_images):
        fichiers = filedialog.askopenfilenames(title="Sélectionnez des images", filetypes=[("Images", "*.png;*.jpg;*.jpeg;*.bmp")])
        for fichier in fichiers:
            listbox_images.insert(tk.END, fichier)
        mettre_a_jour_visionneuse()

    def copier_texte_visionneuse(text_visionneuse):
        texte = text_visionneuse.get("1.0", tk.END)
        pyperclip.copy(texte)
        messagebox.showinfo("Succès", "Texte copié dans le presse-papiers.")

    def generer_document():
        titre = entry_titre.get().strip()
        date_heure = entry_date.get().strip()
        client = entry_client.get().strip()
        version = entry_version.get().strip()
        lecteurs = entry_lecteurs.get().strip()
        sauvegarde = entry_sauvegarde.get().strip()
        postes = listbox_postes.get(0, tk.END)
        poste_serveur = entry_poste_serveur.get().strip()
        description = text_description.get("1.0", tk.END).strip()
        actions = text_actions.get("1.0", tk.END).strip()

        if not titre or not date_heure or not client:
            messagebox.showerror("Erreur", "Le titre, la date/heure, et le client sont obligatoires.")
            return

        doc = Document()
        doc.add_heading(titre, 0)

        doc.add_paragraph(f"Date et Heure : {date_heure}")
        doc.add_paragraph(f"Client : {client}")
        doc.add_paragraph(f"Version installée : {version}")

        doc.add_paragraph(f"Type lecteurs installés : {lecteurs}")
        doc.add_paragraph(f"Marque de lecteur installés : {lecteurs}")
        doc.add_paragraph(f"Fournisseur lecteur : {lecteurs}")
        doc.add_paragraph(f"Présence sauvegarde : {sauvegarde}")
        doc.add_paragraph(f"Nom des postes : {', '.join(postes)}")
        doc.add_paragraph(f"Poste serveur + type poste : {poste_serveur}")

        if description:
            doc.add_heading("Description de l'installation", level=1)
            doc.add_paragraph(description)

        if actions:
            doc.add_heading("Actions à suivre", level=1)
            doc.add_paragraph(actions)

        if listbox_images.size() > 0:
            doc.add_heading("Images", level=1)
            for image_path in listbox_images.get(0, tk.END):
                doc.add_picture(image_path, width=Inches(5.0))

        fichier = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Document Word", "*.docx")])
        if fichier:
            doc.save(fichier)
            messagebox.showinfo("Succès", "Le ticket d'installation a été généré avec succès.")
            os.startfile(fichier)

    btn_generer = tk.Button(scrollable_frame, text="Générer et Ouvrir Document Word", command=generer_document)
    btn_generer.grid(row=14, column=0, columnspan=3, padx=10, pady=20)

