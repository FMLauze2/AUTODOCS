import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os

# Chemin vers le modèle de compte rendu
chemin_modele_compte_rendu = os.path.join(os.path.dirname(__file__), 'COMPTE RENDU.docx')

def create_creation_compte_rendu_tab(tab_creation_compte_rendu):
    # Créer un canvas pour le scroll
    canvas = tk.Canvas(tab_creation_compte_rendu)
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Ajouter un scrollbar vertical lié au canvas
    scrollbar = tk.Scrollbar(tab_creation_compte_rendu, orient=tk.VERTICAL, command=canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    # Créer un frame à l'intérieur du canvas pour y ajouter les widgets
    scrollable_frame = tk.Frame(canvas)

    # Configurer le canvas pour utiliser le scrollbar
    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(
            scrollregion=canvas.bbox("all")
        )
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

    # Widgets du formulaire
    label_titre = tk.Label(scrollable_frame, text="Titre du compte rendu :")
    label_titre.grid(row=0, column=0, padx=10, pady=5, sticky="e")
    entry_titre = tk.Entry(scrollable_frame, width=70)
    entry_titre.grid(row=0, column=1, padx=10, pady=5, sticky="w")

    label_date = tk.Label(scrollable_frame, text="Date et heure :")
    label_date.grid(row=1, column=0, padx=10, pady=5, sticky="e")
    entry_date = tk.Entry(scrollable_frame, width=70)
    entry_date.insert(0, datetime.now().strftime("%Y-%m-%d %H:%M"))
    entry_date.grid(row=1, column=1, padx=10, pady=5, sticky="w")

    label_participants = tk.Label(scrollable_frame, text="Participants :")
    label_participants.grid(row=2, column=0, padx=10, pady=5, sticky="e")
    entry_participants = tk.Entry(scrollable_frame, width=70)
    entry_participants.grid(row=2, column=1, padx=10, pady=5, sticky="w")

    label_infos_contact = tk.Label(scrollable_frame, text="Infos de contact :")
    label_infos_contact.grid(row=3, column=0, padx=10, pady=5, sticky="e")
    entry_infos_contact = tk.Entry(scrollable_frame, width=70)
    entry_infos_contact.grid(row=3, column=1, padx=10, pady=5, sticky="w")

    label_objet_sujet = tk.Label(scrollable_frame, text="Objet / Sujet :")
    label_objet_sujet.grid(row=4, column=0, padx=10, pady=5, sticky="e")
    entry_objet_sujet = tk.Entry(scrollable_frame, width=70)
    entry_objet_sujet.grid(row=4, column=1, padx=10, pady=5, sticky="w")

    label_notes = tk.Label(scrollable_frame, text="Notes / Détails :")
    label_notes.grid(row=5, column=0, padx=10, pady=5, sticky="ne")
    text_notes = tk.Text(scrollable_frame, wrap='word', height=10, width=70)
    text_notes.grid(row=5, column=1, padx=10, pady=5, sticky="w")

    label_actions = tk.Label(scrollable_frame, text="Actions à suivre :")
    label_actions.grid(row=6, column=0, padx=10, pady=5, sticky="ne")
    text_actions = tk.Text(scrollable_frame, wrap='word', height=5, width=70)
    text_actions.grid(row=6, column=1, padx=10, pady=5, sticky="w")

    label_images = tk.Label(scrollable_frame, text="Images (Screenshots) :")
    label_images.grid(row=7, column=0, padx=10, pady=5, sticky="ne")
    listbox_images = tk.Listbox(scrollable_frame, height=5, width=70)
    listbox_images.grid(row=7, column=1, padx=10, pady=5, sticky="w")

    def ajouter_image():
        fichiers = filedialog.askopenfilenames(title="Sélectionnez des images", filetypes=[("Images", "*.png;*.jpg;*.jpeg;*.bmp")])
        for fichier in fichiers:
            listbox_images.insert(tk.END, fichier)

    btn_ajouter_image = tk.Button(scrollable_frame, text="Ajouter Image", command=ajouter_image)
    btn_ajouter_image.grid(row=8, column=1, padx=10, pady=5, sticky="w")

    # Fonction pour remplacer les balises tout en conservant la mise en forme
    def remplacer_balises_dans_run(paragraph, balises_remplacement):
        """Remplacer les balises dans les runs d'un paragraphe sans modifier la mise en forme."""
        for run in paragraph.runs:
            for balise, remplacement in balises_remplacement.items():
                if balise in run.text:
                    run.text = run.text.replace(balise, remplacement)

    # Fonction pour générer le compte rendu à partir du modèle Word
    def generer_compte_rendu():
        titre = entry_titre.get().strip()
        date_heure = entry_date.get().strip()
        participants = entry_participants.get().strip()
        infos_contact = entry_infos_contact.get().strip()
        objet_sujet = entry_objet_sujet.get().strip()
        notes = text_notes.get("1.0", tk.END).strip()
        actions = text_actions.get("1.0", tk.END).strip()
        images = listbox_images.get(0, tk.END)

        if not titre or not date_heure or not participants or not infos_contact:
            messagebox.showerror("Erreur", "Tous les champs obligatoires doivent être remplis.")
            return

        # Vérifier si le fichier modèle existe
        if not os.path.exists(chemin_modele_compte_rendu):
            messagebox.showerror("Erreur", f"Le fichier {chemin_modele_compte_rendu} est introuvable.")
            return

        # Charger le modèle Word
        doc = Document(chemin_modele_compte_rendu)

        # Dictionnaire de balises et de valeurs de remplacement
        balises_remplacement = {
            "[TITRE]": titre,
            "[DATE]": date_heure,
            "[PARTICIPANTS]": participants,
            "[INFOS_CONTACT]": infos_contact,
            "[OBJET_SUJET]": objet_sujet,
            "[NOTES_DETAIL]": notes,
            "[ACTIONS]": actions,
        }

        # Remplacer les balises dans chaque paragraphe tout en conservant la mise en forme
        for paragraphe in doc.paragraphs:
            remplacer_balises_dans_run(paragraphe, balises_remplacement)

        # Rechercher l'emplacement des images dans le modèle et y insérer les images
        for paragraphe in doc.paragraphs:
            if "[LISTE_IMAGES]" in paragraphe.text:
                paragraphe.clear()  # Supprimer la balise
                for image_path in images:
                    doc.add_picture(image_path, width=Inches(5.0))
                break  # On insère les images une fois, puis on arrête la boucle

        # Sauvegarder le fichier généré
        fichier_sortie = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Document Word", "*.docx")])
        if fichier_sortie:
            doc.save(fichier_sortie)
            messagebox.showinfo("Succès", "Le compte rendu a été généré avec succès.")
            os.startfile(fichier_sortie)

    # Bouton pour générer le compte rendu
    btn_generer = tk.Button(scrollable_frame, text="Générer et Ouvrir le Compte Rendu Word", command=generer_compte_rendu)
    btn_generer.grid(row=9, column=1, padx=10, pady=20)
